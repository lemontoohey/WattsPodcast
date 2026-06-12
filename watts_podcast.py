"""
Watts Podcast Generator -- core pipeline.

Converts a PDF, .txt, or .docx document into a podcast-style audio file
spoken in a cloned voice, in the style of Alan Watts.

Stages:
  1. Reference voice extraction (--rip-voice)
  2. Script generation via Claude
  3. Audio generation (TTS + bass engine + humour)
  4. Mastering and assembly
"""

import json
import os
import re
import subprocess

import numpy as np
import soundfile as sf
import torch
from tqdm import tqdm

from bass_engine import BassEngine
from watts_humour import WattsHumour
from llm_client import complete
from socratic_echo import answer_space


# ---------------------------------------------------------------------------
# Compute configuration
# ---------------------------------------------------------------------------

def configure_compute():
    """
    Hybrid GPU/CPU routing for xtts_v2.
    Routes the transformer (heavy) to GPU, the vocoder to CPU when VRAM
    is limited. Prevents GPU OOM on consumer cards while maintaining quality.
    """
    config = {
        "tts_device": "cpu",
        "vocoder_device": "cpu",
        "torch_dtype": torch.float32,
        "gpu_available": False,
        "vram_gb": 0,
        "hybrid_mode": False,
    }

    if torch.cuda.is_available():
        vram = torch.cuda.get_device_properties(0).total_memory / 1e9
        config["vram_gb"] = vram
        config["gpu_available"] = True

        if vram >= 8.0:
            config["tts_device"] = "cuda"
            config["vocoder_device"] = "cuda"
            config["torch_dtype"] = torch.float16
            print(f"[Compute] Full GPU mode ({vram:.1f}GB VRAM) -- float16")

        elif vram >= 4.0:
            config["tts_device"] = "cuda"
            config["vocoder_device"] = "cpu"
            config["torch_dtype"] = torch.float16
            config["hybrid_mode"] = True
            print(f"[Compute] Hybrid mode ({vram:.1f}GB VRAM) -- transformer->GPU, vocoder->CPU")

        elif vram >= 2.0:
            config["tts_device"] = "cuda"
            config["vocoder_device"] = "cpu"
            config["torch_dtype"] = torch.float32
            config["hybrid_mode"] = True
            print(f"[Compute] Low-VRAM hybrid ({vram:.1f}GB) -- with CPU offload")

        else:
            print(f"[Compute] CPU mode (GPU VRAM {vram:.1f}GB insufficient)")
    else:
        print("[Compute] CPU mode (no CUDA detected)")

    torch.backends.cudnn.benchmark = True
    torch.set_num_threads(max(4, torch.get_num_threads()))
    return config


# ---------------------------------------------------------------------------
# TTS model loading with hybrid routing
# ---------------------------------------------------------------------------

def load_tts_model(config):
    from TTS.api import TTS

    tts = TTS(model_name="tts_models/multilingual/multi-dataset/xtts_v2")

    if config["hybrid_mode"]:
        model = tts.synthesizer.tts_model
        if hasattr(model, 'gpt'):
            model.gpt = model.gpt.to(config["tts_device"])
            if config["torch_dtype"] == torch.float16:
                model.gpt = model.gpt.half()

            def to_cpu_hook(module, input, output):
                if isinstance(output, torch.Tensor):
                    return output.to("cpu").float()
                return output

            model.gpt.register_forward_hook(to_cpu_hook)

        if hasattr(model, 'hifigan_decoder'):
            model.hifigan_decoder = model.hifigan_decoder.to("cpu")
    else:
        tts = tts.to(config["tts_device"])

    return tts


def clear_gpu_cache_between_chunks():
    if torch.cuda.is_available():
        torch.cuda.empty_cache()
        torch.cuda.synchronize()


def generate_chunk_with_fallback(tts, text, reference_wav, output_path, config):
    try:
        tts.tts_to_file(text=text, speaker_wav=reference_wav,
                        language="en", file_path=output_path)
    except RuntimeError as e:
        if "out of memory" in str(e).lower():
            print("  [GPU OOM] Falling back to CPU for this chunk...")
            torch.cuda.empty_cache()
            tts.synthesizer.tts_model.to("cpu")
            tts.tts_to_file(text=text, speaker_wav=reference_wav,
                            language="en", file_path=output_path)
            tts.synthesizer.tts_model.to(config["tts_device"])
        else:
            raise


# ---------------------------------------------------------------------------
# Stage 1 -- Reference voice extraction
# ---------------------------------------------------------------------------

def rip_voice(url, start_seconds=90):
    """
    Download audio from a YouTube URL, extract a 45-second segment
    starting at start_seconds, and normalise to -16 LUFS at 22050Hz mono.
    Saves reference_voice.wav.
    """
    print(f"[Stage 1] Ripping voice from YouTube (start={start_seconds}s)...")

    subprocess.run([
        "yt-dlp", "-x", "--audio-format", "wav",
        "-o", "raw_voice.wav", url
    ], check=True)

    subprocess.run([
        "ffmpeg", "-y", "-i", "raw_voice.wav",
        "-ss", str(start_seconds), "-t", "45",
        "-ar", "22050", "-ac", "1",
        "reference_voice_raw.wav"
    ], check=True)

    subprocess.run([
        "ffmpeg-normalize", "reference_voice_raw.wav",
        "-o", "reference_voice.wav",
        "-ar", "22050", "-c:a", "pcm_s16le",
        "--loudness-range-target", "11",
        "-t", "-16"
    ], check=True)

    os.remove("raw_voice.wav")
    os.remove("reference_voice_raw.wav")
    print("[Stage 1] reference_voice.wav saved.")


# ---------------------------------------------------------------------------
# Stage 2 -- Script generation
# ---------------------------------------------------------------------------

def extract_text(path):
    """Extract plain text from PDF, .docx, or .txt input."""
    ext = os.path.splitext(path)[1].lower()
    if ext == '.pdf':
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return '\n'.join(p.extract_text() or '' for p in pdf.pages)
    elif ext == '.docx':
        from docx import Document
        d = Document(path)
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(' | '.join(c.text for c in row.cells))
        return '\n'.join(parts)
    elif ext == '.txt':
        return open(path, encoding='utf-8').read()
    else:
        raise ValueError(f'Unsupported input type: {ext}')


SYSTEM_PROMPT = """
You are rewriting a technical document as a spoken monologue in the voice
and style of Alan Watts.

STYLE RULES:
- Speak directly: 'you see...', 'now consider...', 'do you notice...'
- Short declarative sentences followed by expansive elaboration.
- Replace ALL math notation with verbal descriptions.
- Replace variable names with plain English.
- Conceptual moves: paradox, observer/observed, game metaphor,
  organic analogy, sudden cosmic scale drop.
- Retain ALL scientific accuracy. Only translate the register.
- Target length: {target_words} words.
- Begin mid-thought. Never 'Today we will discuss'.
- No headers. No bullets. Pure flowing prose.

MARKER SYSTEM -- insert these tokens exactly:
  [BREATH]          -- natural pause, every 150-200 words
  [EMPHASIS]text[/EMPHASIS] -- one per paragraph max, sparingly
  [CONTEMPLATE]     -- max 3 per script, genuine philosophical depth
  [LAUGH:TYPE]      -- humour injection point (see humour system)

HUMOUR INTEGRATION:
At appropriate moments insert [LAUGH:TYPE] where TYPE is one of:
  COSMIC_ABSURDITY, SELF_DEPRECATING, ZEN_PUNCHLINE,
  BRITISH_DRY, AUDIENCE_MIRROR
The humour system will replace these tokens with contextual jokes.
Use 2-4 laugh tokens per 20 minutes of content.

CONTEMPLATION QUESTIONS:
At each [CONTEMPLATE] marker, the question generated must:
- Invite the listener to find a NEW USE or APPLICATION for what
  was just explained. Not 'reflect on this' but 'what else could
  this principle govern that nobody has thought of yet?'
- Reference something specific from the preceding passage
- Open genuinely outward -- toward invention, not introspection
- Sound like Watts, not like a homework question

CHAPTER MAP:
After the monologue, output a JSON block in <chapters>...</chapters> tags:
[
  {"title": "The Game Light Plays", "approx_word": 0},
  {"title": "Why the Hollow Sphere", "approx_word": 340}
]
Maximum 8 chapters. Titles must sound like Watts lecture titles.
"""


def generate_script(text, duration_minutes, client, memory_context=''):
    """
    Generate the Watts monologue script via Claude.
    Splits source text into <=12,000-word segments if needed.
    memory_context: optional context block injected into the system prompt
    (used by the Living Memory Thread).
    Saves watts_script.txt and chapters.json.
    """
    target_words = duration_minutes * 140
    words = text.split()
    seg_size = 12000
    segments = [' '.join(words[i:i + seg_size])
                 for i in range(0, len(words), seg_size)] or ['']
    per_segment_words = max(400, target_words // len(segments))

    system = SYSTEM_PROMPT.replace('{target_words}', str(per_segment_words))
    if memory_context:
        system += '\n' + memory_context

    full_script, all_chapters, word_offset = [], [], 0
    for n, seg in enumerate(segments):
        print(f"[Stage 2] Generating script segment {n + 1}/{len(segments)}...")
        out = complete(
            client,
            system=system,
            user=f'Rewrite this as the monologue:\n\n{seg}',
            max_tokens=8000,
        )
        m = re.search(r'<chapters>(.*?)</chapters>', out, re.DOTALL)
        if m:
            try:
                chs = json.loads(m.group(1))
                for c in chs:
                    c['approx_word'] += word_offset
                all_chapters.extend(chs)
            except json.JSONDecodeError:
                pass
            out = re.sub(r'<chapters>.*?</chapters>', '', out, flags=re.DOTALL)
        out = out.strip()
        word_offset += len(out.split())
        full_script.append(out)

    script = '\n\n'.join(full_script)
    open('watts_script.txt', 'w', encoding='utf-8').write(script)
    json.dump(all_chapters[:8], open('chapters.json', 'w'), indent=2)
    print(f"[Stage 2] watts_script.txt saved ({len(script.split())} words).")
    return script


# ---------------------------------------------------------------------------
# Stage 2b -- Script parsing
# ---------------------------------------------------------------------------

def parse_script_markers(script):
    """Split a generated script into ordered audio-generation segments."""
    segments = []
    pattern = (
        r'(\[BREATH\]|\[CONTEMPLATE\]'
        r'|\[EMPHASIS\].*?\[/EMPHASIS\]'
        r'|\[LAUGH:[A-Z_]+\])'
    )
    parts = re.split(pattern, script, flags=re.DOTALL)

    for part in parts:
        part = part.strip()
        if not part:
            continue
        elif part == '[BREATH]':
            segments.append({'type': 'breath', 'duration': 0.8})
        elif part == '[CONTEMPLATE]':
            segments.append({'type': 'contemplate', 'duration': 2.5})
        elif part.startswith('[EMPHASIS]'):
            text = re.sub(r'\[/?EMPHASIS\]', '', part).strip()
            segments.append({'type': 'emphasis', 'text': text})
        elif part.startswith('[LAUGH:'):
            laugh_type = re.search(r'\[LAUGH:([A-Z_]+)\]', part).group(1)
            segments.append({'type': 'laugh', 'laugh_type': laugh_type})
        else:
            # Split into ~250-word TTS-safe chunks at sentence boundaries
            sentences = re.split(r'(?<=[.!?])\s+', part)
            chunk = ''
            for s in sentences:
                if len((chunk + ' ' + s).split()) > 250:
                    if chunk:
                        segments.append({'type': 'text', 'text': chunk.strip()})
                    chunk = s
                else:
                    chunk += ' ' + s
            if chunk.strip():
                segments.append({'type': 'text', 'text': chunk.strip()})
    return segments


# ---------------------------------------------------------------------------
# Stage 3 -- Audio generation
# ---------------------------------------------------------------------------

def calculate_information_density(text):
    words = text.split()
    return sum(1 for w in words if len(w) > 8) / max(len(words), 1)


def apply_emphasis(audio, sr=22050, n_steps=0.3):
    import librosa
    pitched = librosa.effects.pitch_shift(
        audio.astype('float32'), sr=sr, n_steps=n_steps
    )
    return librosa.effects.time_stretch(pitched, rate=0.95)


def generate_silence(duration, sr=22050):
    return np.zeros(int(duration * sr))


def generate_audio(segments, tts, config, humour_system, reflection_fn,
                    bass_engine, arc_director, journal=None, answer_space_s=20.0):
    """
    Render a fully parsed script (see parse_script_markers) to a single
    mono audio array, weaving in breaths, emphasis, contemplation bass
    beds + reflection questions + answer spaces, and humour interjections.

    arc_director: ArcDirector instance. When enabled, drives breath
    duration, contemplate bass parameters, emphasis strength, and
    dense-passage pauses from the script's mapped emotional arc. When
    disabled, fixed v2 defaults are used.

    journal: optional ContemplationJournal -- if provided, each
    contemplation question is logged with its audio timestamp.

    answer_space_s: seconds of thinking-space drone after each
    contemplation question (0 disables it).
    """
    chunks = []
    contemplate_count = 0
    prev_text = ''
    word_count = 0
    total_samples = 0
    sr = 22050

    for i, seg in enumerate(tqdm(segments, desc='[Stage 3] Generating audio')):
        stype = seg['type']

        if stype == 'text':
            density = calculate_information_density(seg['text'])
            path = f'chunk_{i:04d}.wav'

            if density > 0.25:
                # Dense passage: synthesise sentence-by-sentence with
                # extra silence between sentences.
                pause = arc_director.density_pause(word_count)
                sents = seg['text'].split('. ')
                expanded = []
                for s in sents:
                    tmp = f'tmp_{i}_s.wav'
                    generate_chunk_with_fallback(
                        tts, s, 'reference_voice.wav', tmp, config
                    )
                    a, sr = sf.read(tmp)
                    expanded.append(a)
                    expanded.append(generate_silence(pause, sr))
                audio = np.concatenate(expanded)
            else:
                generate_chunk_with_fallback(
                    tts, seg['text'], 'reference_voice.wav', path, config
                )
                audio, sr = sf.read(path)

            chunks.append(audio)
            total_samples += len(audio)
            word_count += len(seg['text'].split())
            prev_text = seg['text']

        elif stype == 'breath':
            duration = arc_director.breath_duration(word_count)
            silence = generate_silence(duration, sr)
            chunks.append(silence)
            total_samples += len(silence)

        elif stype == 'contemplate':
            if contemplate_count < 3:
                p = arc_director.bass_params(word_count)
                dyn_bass = BassEngine(
                    sr=22050,
                    duration=p['duration'],
                    master_level=p['master_level'],
                    glide_time=p['glide_time'],
                    layer3_gain=p['layer3_gain'],
                ) if arc_director.enabled else bass_engine

                tone = dyn_bass.generate_contemplate_layer()
                chunks.append(tone)
                total_samples += len(tone)

                q = reflection_fn(prev_text)
                if journal is not None:
                    journal.log_question(total_samples / sr, q)

                q_path = f'reflect_{contemplate_count}.wav'
                generate_chunk_with_fallback(
                    tts, q, 'reference_voice.wav', q_path, config
                )
                q_audio, _ = sf.read(q_path)

                pre_silence = generate_silence(0.5, sr)
                chunks.append(pre_silence)
                chunks.append(q_audio)
                total_samples += len(pre_silence) + len(q_audio)

                if answer_space_s > 0:
                    intensity = float(arc_director.zone_at(word_count)['intensity'])
                    space = answer_space(BassEngine, seconds=answer_space_s,
                                          sr=sr, intensity=intensity)
                    chunks.append(space)
                    total_samples += len(space)
                else:
                    post_silence = generate_silence(1.0, sr)
                    chunks.append(post_silence)
                    total_samples += len(post_silence)

                contemplate_count += 1

        elif stype == 'emphasis':
            path = f'chunk_{i:04d}.wav'
            generate_chunk_with_fallback(
                tts, seg['text'], 'reference_voice.wav', path, config
            )
            audio, sr = sf.read(path)
            n_steps = arc_director.emphasis_strength(word_count)
            emphasised = apply_emphasis(audio, sr, n_steps=n_steps)
            chunks.append(emphasised)
            total_samples += len(emphasised)
            word_count += len(seg['text'].split())
            prev_text = seg['text']

        elif stype == 'laugh':
            laugh_audio = humour_system.get_laugh_audio(seg['laugh_type'], prev_text)
            if laugh_audio is not None:
                chunks.append(laugh_audio)
                total_samples += len(laugh_audio)

        if i % 3 == 0 and config['hybrid_mode']:
            clear_gpu_cache_between_chunks()

    return np.concatenate(chunks)


# ---------------------------------------------------------------------------
# Stage 4 -- Mastering and assembly
# ---------------------------------------------------------------------------

def master_audio(input_wav, output_mp3):
    filt = ','.join([
        'highpass=f=80',
        'equalizer=f=180:width_type=o:width=1.5:g=1.5',
        'equalizer=f=3000:width_type=o:width=1:g=-1.0',
        'aecho=0.7:0.4:35|45:0.25|0.15',
        'acompressor=threshold=-18dB:ratio=2.5:attack=15:release=200:makeup=2',
        'loudnorm=I=-16:TP=-1.5:LRA=11'
    ])
    subprocess.run([
        'ffmpeg', '-y', '-i', input_wav,
        '-af', filt,
        '-codec:a', 'libmp3lame', '-b:a', '192k',
        '-id3v2_version', '3',
        output_mp3
    ], check=True)


# ---------------------------------------------------------------------------
# Creative contemplation question generation
# ---------------------------------------------------------------------------

CONTEMPLATION_SYSTEM_PROMPT = """
You are generating a spoken contemplation question in the voice of Alan Watts,
to be inserted into a technical monologue after a moment of deep explanation.

YOUR GOAL: Generate a question that drives CREATIVE AND APPLIED THINKING.
Not introspection. Not 'sit with this feeling.'
Instead: 'What could you BUILD, DESIGN, INVENT, or REFRAME using this principle?'

THE QUESTION MUST:
1. Reference something specific and concrete from the preceding passage
2. Suggest a domain or application that is UNEXPECTED --
   not the obvious use of the information.
   Something the inventor of the technique has probably not considered.
   Something that makes the listener think: 'Wait... could that actually work?'
3. Be framed as a genuine possibility, not a rhetorical device.
   The listener should actually want to answer it.
4. Sound like Watts -- unhurried, direct, slightly amused.
5. Be ONE question only. Maximum 3 sentences total.
6. End open -- no answer implied.

QUESTION STRUCTURES THAT WORK:
- 'If [specific mechanism just described] governs [known domain],
  what entirely different domain might it also govern -- and what would
  that mean for how we approach [unexpected application]?'
- 'The [specific phenomenon] happens because [mechanism]. Now:
  where else does that exact logic apply that nobody has looked yet?'
- 'You have just understood why [X] does [Y]. What is the most
  unlike-X thing you can think of that might do Y for the same reason?'

QUESTION STRUCTURES TO AVOID:
- 'Can you notice how this relates to your own experience?'
- 'Sit with that for a moment.'
- 'What does this mean for you personally?'
- Anything that points inward rather than outward.
- Anything that has an obvious answer.

Return only the question text. No preamble. No explanation.
"""


def generate_reflection_question(preceding_text, client):
    """
    Generate a creative-outward contemplation question.
    The question drives invention and novel application,
    not introspection.
    """
    text = complete(
        client,
        system=CONTEMPLATION_SYSTEM_PROMPT,
        user=(
            'Generate a creative contemplation question based on this passage:\n\n'
            f'{preceding_text[-600:]}'
        ),
        max_tokens=120,
    )
    return text.strip()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    import argparse
    import datetime

    from arc_director import ArcDirector
    from llm_client import get_client
    from socratic_echo import ContemplationJournal, generate_reply_episode
    from watts_memory import WattsMemory

    parser = argparse.ArgumentParser(description='Watts Podcast Generator')
    parser.add_argument('--rip-voice', type=str)
    parser.add_argument('--rip-start', type=int, default=90)
    parser.add_argument('--rip-laugh', type=str)
    parser.add_argument('--laugh-start', type=int, default=0)
    parser.add_argument('--process-laugh', action='store_true')
    parser.add_argument('--input', type=str)
    parser.add_argument('--duration', type=int, default=20)
    parser.add_argument('--use-existing-script', action='store_true')
    parser.add_argument('--compute-info', action='store_true')
    # --- Dynamic Editions ---
    parser.add_argument('--no-memory', action='store_true',
                         help='standalone episode, no memory context injected')
    parser.add_argument('--memory-info', action='store_true')
    parser.add_argument('--forget', type=int,
                         help='delete episode N from memory')
    parser.add_argument('--no-arc', action='store_true',
                         help='fixed v2 render parameters')
    parser.add_argument('--answer-space', type=float, default=20.0,
                         help='thinking-space seconds after each question; 0=off')
    parser.add_argument('--reply', type=str,
                         help='path to a filled journal_ep<N>.md')
    args = parser.parse_args()

    config = configure_compute()
    memory = WattsMemory(listener='Liam')

    if args.compute_info:
        print(config)
        return
    if args.memory_info:
        memory.info()
        return
    if args.forget:
        memory.data['episodes'] = [e for e in memory.data['episodes']
                                    if e['id'] != args.forget]
        memory.save()
        print(f'[Memory] Episode {args.forget} forgotten.')
        return

    if args.rip_laugh:
        from watts_humour import rip_laugh
        rip_laugh(args.rip_laugh, args.laugh_start)
        return

    if args.process_laugh:
        from watts_humour import process_laugh
        process_laugh()
        return

    if args.rip_voice:
        rip_voice(args.rip_voice, args.rip_start)

    # ---------- REPLY MODE ----------
    if args.reply:
        client = get_client()
        script_path, ep_id = generate_reply_episode(args.reply, client, memory)
        if script_path is None:
            return

        tts = load_tts_model(config)
        bass = BassEngine(sr=22050)
        humour = WattsHumour(client, tts, 'reference_voice.wav', config)

        script = open(script_path, encoding='utf-8').read()
        arc = ArcDirector(script, client, enabled=not args.no_arc)
        segments = parse_script_markers(script)

        audio = generate_audio(
            segments, tts, config, humour,
            lambda t: generate_reflection_question(t, client),
            bass, arc, journal=None, answer_space_s=args.answer_space
        )

        wav_out = f'watts_reply_ep{ep_id}.wav'
        mp3_out = f'watts_reply_ep{ep_id}.mp3'
        sf.write(wav_out, audio, 22050)
        master_audio(wav_out, mp3_out)
        print(f'\n[Done] Reply episode: {mp3_out}')
        return

    # ---------- NORMAL EPISODE ----------
    if args.input:
        client = get_client()
        tts = load_tts_model(config)

        bass = BassEngine(sr=22050, duration=12.0, master_level=0.35)
        humour = WattsHumour(client, tts, 'reference_voice.wav', config)

        if not args.use_existing_script:
            text = extract_text(args.input)
            mem_ctx = '' if args.no_memory else memory.memory_context()
            generate_script(text, args.duration, client, memory_context=mem_ctx)

        script = open('watts_script.txt', encoding='utf-8').read()
        arc = ArcDirector(script, client, enabled=not args.no_arc)
        segments = parse_script_markers(script)

        ep_id = len(memory.data['episodes']) + 1
        journal = ContemplationJournal(ep_id, 'pending', os.path.basename(args.input))

        audio = generate_audio(
            segments, tts, config, humour,
            lambda t: generate_reflection_question(t, client),
            bass, arc, journal=journal, answer_space_s=args.answer_space
        )

        ts = datetime.datetime.now().strftime('%Y%m%d_%H%M')
        base = os.path.splitext(args.input)[0]
        wav_out = f'{base}_watts_{ts}.wav'
        mp3_out = f'{base}_watts_{ts}.mp3'
        sf.write(wav_out, audio, 22050)
        master_audio(wav_out, mp3_out)

        ep = memory.remember_episode(script, os.path.basename(args.input), client)
        journal.episode_title = ep['title']
        journal.write()

        print(f'\n[Done] {mp3_out}')
        print(f'[Done] Journal: journal_ep{ep["id"]}.md -- answer it, then run '
              f'--reply journal_ep{ep["id"]}.md')


if __name__ == '__main__':
    main()
